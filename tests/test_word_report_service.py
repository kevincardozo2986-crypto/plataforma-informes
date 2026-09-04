import pandas as pd
from docx import Document
from openpyxl import load_workbook

from app.services.excel_service import ExcelProcess
from app.services.word_report_service import _bundled_template_path, generate_word_report


def test_institutional_word_template_is_part_of_project():
    template = _bundled_template_path()
    assert template.is_file()
    assert template.name == "PLANTILLA_INFORME.docx"


def test_generates_word_from_finished_excel_and_preserves_template(tmp_path):
    original = pd.DataFrame(
        {
            "curso": ["Curso A", "Curso A", "Curso A", "Curso B", "Curso B", "Curso A"],
            "idusuario": [1, 1, 10, 11, 11, 10],
            "usuario": ["Docente Uno", "Docente Uno", "Ana", "Beto", "Beto", "Ana"],
            "rol": ["editingteacher", "editingteacher", "student", "student", "student", "student"],
            "Mes": [2, 3, 2, 2, 3, 3],
            "Dia": [1, 2, 3, 4, 5, 6],
            "accion": ["updated", "viewed", "viewed", "created", "viewed", "submitted"],
        }
    )
    process = ExcelProcess(tmp_path / "Informe.xlsx")
    process.create_original_from_chunks([original])
    process.crear_tabla_docentes()
    process.crear_grafica_docentes("Ingenieria de Sistemas", "2026-1")
    process.crear_tabla_estudiantes("Ingenieria de Sistemas", "2026-1")
    process.crear_grafica_estudiantes("Ingenieria de Sistemas", "2026-1")
    process.crear_tabla_actividades("Ingenieria de Sistemas", "2026-1")
    process.crear_diseno_cursos("Ingenieria de Sistemas", "2026-1")

    workbook = load_workbook(process.path, read_only=False, data_only=True)
    assert "Resumen Informe" in workbook.sheetnames
    assert len(workbook["Resumen Informe"]._charts) == 2
    workbook.close()

    destination = tmp_path / "Informe_2026-1_ING_SIS.docx"
    result = generate_word_report(
        process.path, destination, "Administracion de Empresas", "2027-2"
    )

    assert result == destination
    assert destination.is_file()
    document = Document(destination)
    assert len(document.tables) == 5
    assert len(document.inline_shapes) == 6
    assert document.tables[0].cell(1, 1).text == "6"
    assert document.tables[2].cell(1, 0).text == "Febrero"
    report_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    xml_text = " ".join(document.element.itertext())
    normalized_text = report_text.casefold()
    assert "Administracion de Empresas" in report_text
    assert "2027-2" in report_text
    assert "andrea cruz yomayusa" not in normalized_text
    assert "optimizacion_6a_6b" not in normalized_text
    assert "400.568" not in report_text
    assert "59 cursos" not in report_text
    assert "Días al mes de uso de la plataforma Open LMS - Docentes" in report_text
    assert "Cursos con mayor continuidad estudiantil" in report_text
    assert "Administracion de Empresas" in xml_text
    assert "2027-2" in xml_text
    assert "10 hojas" not in report_text
    assert "Hoja «Docentes»" not in report_text
    assert "Hoja «Estudiantes»" not in report_text
    assert "9 hojas" in report_text
    assert "{{" not in xml_text
    caption = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.startswith("Ilustración 2 ")
    )
    assert caption.text == (
        "Ilustración 2 Días al mes de uso de la plataforma Open LMS "
        "- Docentes 2027-2"
    )
    assert all(run.italic for run in caption.runs)
    assert all(run.font.size.pt == 9 for run in caption.runs)
    assert caption.style.name == "Caption"
    field_codes = " ".join(
        node.text or "" for node in document.element.iter()
        if node.tag.endswith("}instrText")
    )
    assert 'TOC \\h \\z \\c "FiguraInforme"' in field_codes
    assert field_codes.count("SEQ FiguraInforme") == 6
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    assert "Tabla de ilustraciones" in paragraphs
    assert "Tabla de contenido" in paragraphs
    assert "PAGEREF" not in field_codes
    toc_cache = next(p for p in paragraphs if "Resumen Ejecutivo" in p and "1. Introducción" in p)
    assert "7. Diseño de cursos virtuales" in toc_cache
    assert "TOC \\h \\z \\u" in field_codes
    assert "1. Introducción" in paragraphs
    assert "7. Diseño de cursos virtuales" in paragraphs


def test_toc_field_matches_illustrations_structure():
    template = _bundled_template_path()
    document = Document(template)
    paragraphs = list(document.paragraphs)

    def _field_layout(paragraph):
        return [
            (node.get("{%s}fldCharType" % node.nsmap.get("w")), node.get("{%s}dirty" % node.nsmap.get("w")))
            for node in paragraph._p.iter()
            if node.tag.endswith("}fldChar")
        ]

    toc = next(p for p in paragraphs if "La tabla de contenido se actualizar" in p.text)
    tof = next(p for p in paragraphs if "La tabla de ilustraciones se actualizar" in p.text)
    assert [kind for kind, _ in _field_layout(toc)] == [kind for kind, _ in _field_layout(tof)]
    assert _field_layout(toc)[0] == ("begin", "true")
    assert _field_layout(tof)[0] == ("begin", "true")
    toc_codes = " ".join(
        node.text or "" for node in toc._p.iter() if node.tag.endswith("}instrText")
    )
    tof_codes = " ".join(
        node.text or "" for node in tof._p.iter() if node.tag.endswith("}instrText")
    )
    assert toc_codes.strip() == r"TOC \h \z \u"
    assert tof_codes.strip() == r'TOC \h \z \c "FiguraInforme"'
