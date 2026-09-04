"""Generacion del informe Word institucional desde el Excel terminado."""

import re
import sys
from pathlib import Path
from tempfile import TemporaryDirectory

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
from openpyxl import load_workbook


TEMPLATE_FILENAME = "PLANTILLA_INFORME.docx"


def _bundled_template_path():
    """Localiza la plantilla tanto en el proyecto como en una app empaquetada."""
    bundle_root = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parents[2]))
    candidates = (
        bundle_root / "templates" / TEMPLATE_FILENAME,
        Path(sys.executable).resolve().parent / "templates" / TEMPLATE_FILENAME,
        Path(__file__).resolve().parents[2] / "templates" / TEMPLATE_FILENAME,
        Path(__file__).resolve().parents[2] / TEMPLATE_FILENAME,
    )
    return next((path for path in candidates if path.is_file()), candidates[0])


TEMPLATE_PATH = _bundled_template_path()
MONTH_NAMES = {
    "ENE": "Enero", "FEB": "Febrero", "MAR": "Marzo", "ABR": "Abril",
    "MAY": "Mayo", "JUN": "Junio", "JUL": "Julio", "AGO": "Agosto",
    "SEP": "Septiembre", "OCT": "Octubre", "NOV": "Noviembre", "DIC": "Diciembre",
}
BLUE = "#2878B5"
GOLD = "#D6A419"
NAVY = "#17365D"


def _number(value, decimals=0):
    if value in (None, ""):
        return "0"
    value = float(value)
    if decimals:
        return f"{value:,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{int(round(value)):,}".replace(",", ".")


def _find_row(sheet, label):
    target = label.strip().casefold()
    for row in range(1, sheet.max_row + 1):
        if str(sheet.cell(row, 1).value or "").strip().casefold() == target:
            return row
    raise ValueError(f"No se encontro el bloque '{label}' en Resumen Informe.")


def _read_block(sheet, label, columns):
    row = _find_row(sheet, label) + 2
    rows = []
    while row <= sheet.max_row and sheet.cell(row, 1).value not in (None, ""):
        rows.append(tuple(sheet.cell(row, column).value for column in range(1, columns + 1)))
        row += 1
    return rows


def _extract_data(workbook_path):
    workbook = load_workbook(workbook_path, read_only=True, data_only=True)
    try:
        required = {"Resumen Informe", "Docentes DG", "Tabla Dinamica Estudiantes", "Dise\u00f1o de Cursos"}
        missing = required.difference(workbook.sheetnames)
        if missing:
            raise ValueError("El Excel no contiene las hojas requeridas: " + ", ".join(sorted(missing)))

        summary = workbook["Resumen Informe"]
        indicators = _read_block(summary, "Indicadores principales", 3)
        monthly = _read_block(summary, "Actividad mensual", 4)
        courses = _read_block(summary, "Cursos destacados", 4)
        courses.sort(
            key=lambda row: (
                -float(row[3] or 0), -float(row[1] or 0), str(row[0]).casefold()
            )
        )
        teachers = _read_block(summary, "Continuidad docente", 3)

        teacher_chart = []
        for row in workbook["Docentes DG"].iter_rows(min_row=2, values_only=True):
            if row[0] not in (None, "") and isinstance(row[1], (int, float)):
                teacher_chart.append((str(row[0]), float(row[1])))

        students_sheet = workbook["Tabla Dinamica Estudiantes"]
        headers = [str(cell.value or "").strip().upper() for cell in students_sheet[3]]
        total_days_column = headers.index("TOTAL") + 1
        student_start_column = total_days_column + 1
        total_students_column = headers.index("TOTAL", total_days_column) + 1
        total_row = next(
            row for row in range(4, students_sheet.max_row + 1)
            if str(students_sheet.cell(row, 1).value or "").strip().upper() == "TOTAL GENERAL"
        )
        student_days = [
            (headers[column - 1], float(students_sheet.cell(total_row, column).value or 0))
            for column in range(2, total_days_column)
        ]
        student_users = [
            (headers[column - 1], float(students_sheet.cell(total_row, column).value or 0))
            for column in range(student_start_column, total_students_column)
        ]

        design_sheet = workbook["Dise\u00f1o de Cursos"]
        design = {
            str(design_sheet.cell(5, column).value or ""): int(design_sheet.cell(6, column).value or 0)
            for column in range(3, 7)
        }
        return {
            "sheet_names": list(workbook.sheetnames),
            "indicators": indicators,
            "monthly": monthly,
            "courses": courses,
            "teachers": teachers,
            "teacher_chart": teacher_chart,
            "student_days": student_days,
            "student_users": student_users,
            "design": design,
        }
    finally:
        workbook.close()


def _line_chart(path, categories, series, title, ylabel):
    figure, axis = plt.subplots(figsize=(8, 5.4), dpi=240)
    figure.patch.set_facecolor('white')
    axis.set_facecolor('#FFFFFF')
    
    for name, values, color in series:
        axis.plot(categories, values, "o-", linewidth=2.8, markersize=8, label=name, color=color)
        for category, value in zip(categories, values):
            axis.annotate(
                _number(value, 1 if isinstance(value, float) and not value.is_integer() else 0),
                (category, value), xytext=(0, 10), textcoords="offset points",
                ha="center", fontsize=8, fontweight="bold", color=color,
            )
    
    axis.set_title(title, color=NAVY, fontweight="bold", fontsize=14, pad=16)
    axis.set_ylabel(ylabel, fontsize=11, color=NAVY, fontweight="600")
    axis.set_xlabel("", fontsize=11)
    axis.tick_params(axis='y', labelsize=10, colors=NAVY)
    axis.tick_params(axis='x', labelsize=10, colors=NAVY)
    axis.grid(axis="y", alpha=0.15, linestyle="--", linewidth=0.7)
    
    if len(series) > 1:
        axis.legend(frameon=True, fancybox=True, shadow=False, fontsize=10, 
                   loc='upper left', framealpha=0.95, edgecolor='#E3E6ED')
    
    axis.spines["top"].set_visible(False)
    axis.spines["right"].set_visible(False)
    axis.spines["left"].set_color('#E3E6ED')
    axis.spines["bottom"].set_color('#E3E6ED')
    
    figure.tight_layout()
    figure.savefig(path, dpi=300, bbox_inches="tight", facecolor='white')
    plt.close(figure)


def _bar_chart(path, courses, title):
    names = [str(row[0]).replace("_", " ").title() for row in courses]
    values = [float(row[3] or 0) for row in courses]
    figure, axis = plt.subplots(figsize=(8, 5.4), dpi=240)
    figure.patch.set_facecolor('white')
    axis.set_facecolor('#FFFFFF')
    
    bars = axis.barh(names[::-1], values[::-1], color="#36BCE8", edgecolor="#0A7FA8", linewidth=1.2)
    axis.set_title(title, color=NAVY, fontweight="bold", fontsize=14, pad=16)
    
    # Agregar valores al lado de las barras
    for bar, value in zip(bars, values[::-1]):
        width = bar.get_width()
        axis.text(width + max(values or [1]) * 0.015, bar.get_y() + bar.get_height() / 2,
                  _number(value), va="center", fontsize=10, fontweight="600", color=NAVY)
    
    axis.set_xlabel("Días activos", fontsize=11, fontweight="600", color=NAVY)
    axis.tick_params(axis='x', labelsize=10, colors=NAVY)
    axis.tick_params(axis='y', labelsize=10, colors=NAVY)
    axis.grid(axis="x", alpha=0.15, linestyle="--", linewidth=0.7)
    
    axis.spines["top"].set_visible(False)
    axis.spines["right"].set_visible(False)
    axis.spines["left"].set_color('#E3E6ED')
    axis.spines["bottom"].set_color('#E3E6ED')
    
    figure.tight_layout()
    figure.savefig(path, dpi=300, bbox_inches="tight", facecolor='white')
    plt.close(figure)


def _pie_chart(path, design, title):
    values = [design.get("Sin contenido", 0), design.get("Con contenido", 0)]
    labels = ["Sin contenido", "Con contenido"]
    colors = [GOLD, BLUE]
    figure, axis = plt.subplots(figsize=(8, 5.4), dpi=240)
    figure.patch.set_facecolor('white')

    if sum(values):
        visible = [(label, value, color) for label, value, color in zip(labels, values, colors) if value]
        chart_labels, chart_values, chart_colors = zip(*visible)
        wedges, _texts, autotexts = axis.pie(
            chart_values,
            autopct="%1.1f%%",
            colors=chart_colors,
            startangle=90,
            counterclock=False,
            textprops={"weight": "bold", "fontsize": 11},
            wedgeprops={"edgecolor": "white", "linewidth": 3},
            pctdistance=0.68,
        )
        for autotext in autotexts:
            autotext.set_color("white")
            autotext.set_fontsize(12)
            autotext.set_fontweight("bold")
        axis.legend(
            wedges, chart_labels, loc="lower center", bbox_to_anchor=(0.5, -0.08),
            ncol=len(chart_labels), frameon=False, fontsize=11,
        )
    else:
        axis.text(0.5, 0.5, "Sin datos de diseño", ha="center", va="center", 
                 color=NAVY, fontsize=13, fontweight="bold")

    axis.set_aspect("equal", adjustable="box")
    axis.set_title(title, color=NAVY, fontweight="bold", fontsize=14, pad=16)
    figure.tight_layout()
    figure.savefig(path, dpi=300, bbox_inches="tight", facecolor='white')
    plt.close(figure)


def _create_charts(directory, data, program, period):
    paths = [Path(directory) / f"chart_{index}.png" for index in range(6)]
    monthly = data["monthly"]
    categories = [MONTH_NAMES.get(str(row[0]).upper(), str(row[0])) for row in monthly]
    _line_chart(paths[0], categories, [
        ("Estudiantes", [row[1] or 0 for row in monthly], BLUE),
        ("Docentes", [row[3] or 0 for row in monthly], GOLD),
    ], "Eventos mensuales por tipo de usuario", "Eventos registrados")
    teacher = data["teacher_chart"]
    _line_chart(paths[1], [MONTH_NAMES.get(row[0].upper(), row[0]) for row in teacher],
                [("PROMEDIO", [row[1] for row in teacher], BLUE)],
                "D\u00edas al mes de uso del Campus Virtual por parte de los docentes\n"
                f"de la facultad de {program} {period}", "Promedio de d\u00edas de uso")
    days = data["student_days"]
    _line_chart(paths[2], [MONTH_NAMES.get(row[0].upper(), row[0]) for row in days],
                [("Promedio de días de uso", [row[1] for row in days], BLUE)],
                "D\u00edas del mes de uso del Campus Virtual por parte de los estudiantes\n"
                f"de la facultad de {program} {period}", "Promedio de d\u00edas de uso")
    users = data["student_users"]
    _line_chart(paths[3], [MONTH_NAMES.get(row[0].upper(), row[0]) for row in users],
                [("Promedio de estudiantes", [row[1] for row in users], GOLD)],
                "Promedio de estudiantes que usaron el Campus Virtual de la facultad\n"
                f"de {program} {period}", "Promedio de estudiantes")
    _bar_chart(paths[4], data["courses"], "Cursos con mayor continuidad estudiantil")
    _pie_chart(paths[5], data["design"], "Diseño de Cursos")
    return paths


def _fill_table(table, headers, rows):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    values = [headers, *rows]
    for row_index, table_row in enumerate(table.rows):
        row_values = values[row_index] if row_index < len(values) else [""] * len(table.columns)
        for column_index, cell in enumerate(table_row.cells):
            cell.text = str(row_values[column_index] if column_index < len(row_values) else "")
            
            # Dar estilo a la fila de encabezados
            if row_index == 0:
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco
                # Color de fondo azul para encabezado
                from docx.oxml import parse_xml
                shading_elm = parse_xml(r'<w:shd {} w:fill="2878B5"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                # Dar estilo a las filas de datos
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(33, 21, 104)  # Texto oscuro


def _replace_images(document, image_paths):
    if len(document.inline_shapes) < len(image_paths):
        raise ValueError("La plantilla no contiene los seis espacios de graficas esperados.")
    original_shapes = list(document.inline_shapes)[:len(image_paths)]
    for shape, image_path in zip(original_shapes, image_paths):
        old_drawing = shape._inline.getparent()
        run = old_drawing.getparent()
        with Image.open(image_path) as source_image:
            image_ratio = source_image.width / source_image.height
        max_width, max_height = int(shape.width), int(shape.height)
        width = min(max_width, int(max_height * image_ratio))
        height = int(width / image_ratio)
        temporary_paragraph = document.add_paragraph()
        temporary_run = temporary_paragraph.add_run()
        new_shape = temporary_run.add_picture(
            str(image_path), width=width, height=height
        )
        new_drawing = new_shape._inline.getparent()
        run.replace(old_drawing, new_drawing)
        temporary_paragraph._element.getparent().remove(temporary_paragraph._element)


def _replace_paragraph(document, prefix, text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            paragraph.text = text
            return True
    return False


def _replace_paragraph_group(document, prefix, texts):
    """Actualiza un bloque narrativo sin cambiar su lugar en la plantilla."""
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            for offset, text in enumerate(texts):
                if index + offset < len(document.paragraphs):
                    document.paragraphs[index + offset].text = text
            return True
    return False


def _replace_containing(document, fragment, text):
    for paragraph in document.paragraphs:
        if fragment.casefold() in paragraph.text.casefold():
            paragraph.text = text
            return True
    return False


def _replace_marker(document, marker, text):
    """Reemplaza un marcador tanto en párrafos normales como en cuadros de texto."""
    replaced = False
    for paragraph in document.paragraphs:
        if marker in paragraph.text:
            paragraph.text = paragraph.text.replace(marker, text)
            replaced = True
    for node in document.element.iter(qn("w:t")):
        if marker in (node.text or ""):
            node.text = node.text.replace(marker, text)
            replaced = True
    return replaced


def _replace_cover_fields(document, program, period):
    """Actualiza textos de portada que Word guarda dentro de cuadros de texto."""
    for node in document.element.iter(qn("w:t")):
        original = node.text or ""
        updated = original.replace("{{PROGRAMA}}", program).replace("{{PERIODO}}", period)
        updated = re.sub(
            r"Ingenier.a de Sistemas", program, original,
            flags=re.IGNORECASE,
        ) if updated == original else updated
        updated = updated.replace("2026-1", period)
        if updated != original:
            node.text = updated


def _style_figure_captions(document):
    """Aplica un pie de figura discreto y diferente al texto del informe."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    for paragraph in document.paragraphs:
        if not paragraph.text.strip().startswith("Ilustración "):
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.keep_with_next = False
        for run in paragraph.runs:
            run.italic = True
            run.bold = False
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(84, 101, 120)


def _set_figure_caption(document, index, description):
    """Crea un pie con numeración SEQ para que Word construya la tabla de figuras."""
    paragraph = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(f"Ilustración {index}")
    )
    paragraph.clear()
    paragraph.style = document.styles["Caption"]
    paragraph.add_run("Ilustración ")

    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " SEQ FiguraInforme \\* ARABIC "
    instruction_run._r.append(instruction)

    separator_run = paragraph.add_run()
    separator = OxmlElement("w:fldChar")
    separator.set(qn("w:fldCharType"), "separate")
    separator_run._r.append(separator)
    separator_run.add_text(str(index))

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    paragraph.add_run(f" {description}")


def _configure_illustrations_field(document):
    for node in document.element.iter(qn("w:instrText")):
        if "TOC" in (node.text or "") and ("Caption,1" in node.text or "FiguraInforme" in node.text):
            node.text = ' TOC \\h \\z \\c "FiguraInforme" '


SECTION_TITLES = (
    "Resumen Ejecutivo",
    "1. Introducción",
    "2. Metodología de los indicadores",
    "3. Comportamiento mensual de la actividad",
    "4. Docentes",
    "5. Estudiantes",
    "6. Cursos destacados",
    "7. Diseño de cursos virtuales",
)


def _set_field_cache_lines(field_paragraph_element, lines):
    """Inserta un caché estático entre separate/end para LibreOffice.

    Word regenera el campo al abrir (begin dirty); LibreOffice muestra
    este caché en el PDF porque no ejecuta campos TOC/SEQ de Word.
    """
    runs = list(field_paragraph_element.findall(qn("w:r")))
    sep_idx = end_idx = None
    for i, run in enumerate(runs):
        types = {c.get(qn("w:fldCharType")) for c in run.iter(qn("w:fldChar"))}
        if "separate" in types and sep_idx is None:
            sep_idx = i
        if "end" in types:
            end_idx = i
    if sep_idx is None or end_idx is None or end_idx <= sep_idx + 1:
        # Sin caché previo con formato esperado; si solo hay placeholder,
        # sep_idx+1 == end_idx-? lo reconstruimos igual si hay hueco.
        if sep_idx is None or end_idx is None:
            return False
    for run in runs[sep_idx + 1:end_idx]:
        field_paragraph_element.remove(run)
    anchor = runs[sep_idx]
    for li, line in enumerate(lines):
        cache_run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.set(qn("xml:space"), "preserve")
        text_node.text = line
        cache_run.append(text_node)
        anchor.addnext(cache_run)
        anchor = cache_run
        if li < len(lines) - 1:
            br_run = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br_run.append(br)
            anchor.addnext(br_run)
            anchor = br_run
    return True


def _field_paragraph(document, predicate):
    for paragraph in document.element.iter(qn("w:p")):
        for node in paragraph.iter(qn("w:instrText")):
            if predicate(node.text or ""):
                return paragraph
    return None


def _ensure_toc_cache(document):
    field_p = _field_paragraph(
        document,
        lambda t: "TOC" in t and "FiguraInforme" not in t
        and "PAGEREF" not in t and "SEQ" not in t,
    )
    if field_p is None:
        return False
    return _set_field_cache_lines(field_p, list(SECTION_TITLES))


def _ensure_illustrations_cache(document, figure_captions):
    field_p = _field_paragraph(
        document,
        lambda t: "TOC" in t and "FiguraInforme" in t,
    )
    if field_p is None:
        return False
    lines = [f"Ilustración {i} {caption}" for i, caption in enumerate(figure_captions, 1)]
    return _set_field_cache_lines(field_p, lines)


def _fix_cover_overflow(document, program=""):
    """Evita que la portada recorte el último nombre en Mac/LibreOffice.

    La caja usa <a:noAutofit/> con altura fija (140.9pt) y texto vertical
    btLr a 14pt. Con otra métrica de fuente (Mac) el programa largo hace
    wrap y expulsa la última línea. Se activa auto-ajuste, se amplía el
    fallback VML y se reduce la fuente si el programa es largo.
    """
    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    fixed = 0
    for no_autofit in list(document.element.iter(f"{{{ns_a}}}noAutofit")):
        parent = no_autofit.getparent()
        if parent is None:
            continue
        idx = list(parent).index(no_autofit)
        replacement = OxmlElement("a:spAutoFit")
        parent.remove(no_autofit)
        parent.insert(idx, replacement)
        fixed += 1
    for rect in document.element.iter("{urn:schemas-microsoft-com:vml}rect"):
        style = rect.get("style") or ""
        if "height:140.9pt" in style:
            rect.set("style", style.replace("height:140.9pt", "height:168pt"))
            fixed += 1
    if program and len(program.strip()) > 24:
        for node in document.element.iter(qn("w:t")):
            if (node.text or "").strip().casefold() == program.strip().casefold():
                run = node.getparent()
                rpr = run.find(qn("w:rPr")) if run is not None else None
                if rpr is not None:
                    for sz in rpr.iter(qn("w:sz")):
                        try:
                            if int(sz.get(qn("w:val"), "28")) > 24:
                                sz.set(qn("w:val"), "24")
                        except ValueError:
                            pass
    return fixed


def _build_auto_field_paragraph(instruction, placeholder):
    """Crea un párrafo con un campo auto-actualizable (begin dirty, separate, end)."""
    paragraph_element = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "Normal")
    properties.append(style)
    paragraph_element.append(properties)

    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run.append(begin)
    paragraph_element.append(begin_run)

    instruction_run = OxmlElement("w:r")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    instruction_run.append(instruction_node)
    paragraph_element.append(instruction_run)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    paragraph_element.append(separate_run)

    text_run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = placeholder
    text_run.append(text_node)
    paragraph_element.append(text_run)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph_element.append(end_run)
    return paragraph_element


def _is_uniform_auto_field(element, instruction):
    """Indica si el elemento ya es un campo auto-actualizable con ese código."""
    instructions = [node.text or "" for node in element.iter(qn("w:instrText"))]
    if not any(text.strip() == instruction.strip() for text in instructions):
        return False
    begins = [node for node in element.iter(qn("w:fldChar"))
              if node.get(qn("w:fldCharType")) == "begin"]
    return any(begin.get(qn("w:dirty")) for begin in begins)


def _configure_toc_field(document):
    """Deja la tabla de contenido con la misma estructura que la de ilustraciones.

    Reemplaza el bloque de entradas fijas (control SDT o párrafos con PAGEREF
    y números de página desactualizados) por un único párrafo con campo
    auto-actualizable y marcador begin dirty, para que Word regenere ambas
    tablas al abrir el documento.
    """
    body = document.element.find(qn("w:body"))
    if body is None:
        return False
    children = list(body)

    def _is_toc_instruction(text):
        return ("TOC" in text and "FiguraInforme" not in text
                and "PAGEREF" not in text and "SEQ" not in text)

    start = next(
        (index for index, child in enumerate(children)
         if any(_is_toc_instruction(node.text or "")
                for node in child.iter(qn("w:instrText")))),
        None,
    )
    if start is None:
        return False

    depth = 0
    started = False
    end = None
    for index in range(start, len(children)):
        for field_char in children[index].iter(qn("w:fldChar")):
            char_type = field_char.get(qn("w:fldCharType"))
            if char_type == "begin":
                depth += 1
                started = True
            elif char_type == "end" and started:
                depth -= 1
                if depth == 0:
                    end = index
                    break
        if end is not None:
            break
    if end is None:
        return False

    if start == end and _is_uniform_auto_field(children[start], "TOC \\h \\z \\u"):
        return True
    field_paragraph = _build_auto_field_paragraph(
        " TOC \\h \\z \\u ",
        "La tabla de contenido se actualizará al abrir el documento.",
    )
    children[start].addprevious(field_paragraph)
    for child in children[start:end + 1]:
        body.remove(child)
    return True


def generate_word_report(workbook_path, output_path, program, period, template_path=None):
    """Genera el DOCX final conservando estructura, estilos y posiciones de la plantilla."""
    workbook_path = Path(workbook_path)
    output_path = Path(output_path)
    template_path = Path(template_path) if template_path else _bundled_template_path()
    if not workbook_path.is_file():
        raise FileNotFoundError("El Excel terminado no existe.")
    if not template_path.is_file():
        raise FileNotFoundError(
            "No se encontro la plantilla institucional de Word. "
            f"Debe existir en: {template_path}"
        )

    data = _extract_data(workbook_path)
    indicators = {str(row[0]): row[1:] for row in data["indicators"]}
    total_events = float(indicators.get("Eventos totales", (0, ""))[0] or 0)
    design_total = data["design"].get("Total", 0)
    with_content = data["design"].get("Con contenido", 0)
    content_percent = (with_content * 100 / design_total) if design_total else 0

    with TemporaryDirectory(prefix="informe_word_") as temporary:
        images = _create_charts(temporary, data, program, period)
        document = Document(template_path)
        _replace_images(document, images)
        _replace_cover_fields(document, program, period)
        _fix_cover_overflow(document, program)
        _configure_toc_field(document)
        _configure_illustrations_field(document)
        _ensure_toc_cache(document)

        indicator_rows = []
        for label, result, reading in data["indicators"]:
            decimals = 1 if isinstance(result, float) and not result.is_integer() else 0
            suffix = " promedio" if str(label).startswith("D\u00edas activos") else (
                " eventos" if str(label).startswith("Actividad") else ""
            )
            indicator_rows.append((label, _number(result, decimals) + suffix, reading))
        monthly_rows = [
            (MONTH_NAMES.get(str(row[0]).upper(), row[0]), _number(row[1]), _number(row[2]), _number(row[3]))
            for row in data["monthly"]
        ]
        teacher_rows = [(row[0], row[1], _number(row[2])) for row in data["teachers"]]
        course_rows = [(row[0], _number(row[1]), _number(row[2]), _number(row[3])) for row in data["courses"]]
        _fill_table(document.tables[0], ("Indicador", "Resultado", "Lectura"), indicator_rows)
        _fill_table(document.tables[2], ("Mes", "Eventos estudiantes", "Estudiantes activos", "Eventos docentes"), monthly_rows)
        _fill_table(document.tables[3], ("Curso", "Docente", "D\u00edas activos"), teacher_rows)
        _fill_table(document.tables[4], ("Curso", "Eventos", "Estudiantes \u00fanicos", "D\u00edas activos"), course_rows)

        _replace_paragraph(document, "Durante el semestre", (
            f"Durante el periodo {period} se analizaron {_number(total_events)} eventos de Open LMS "
            f"correspondientes al programa de {program}. Los resultados permiten reconocer el nivel de "
            "participacion de estudiantes y docentes, la continuidad mensual y los cursos con mayor actividad."
        ))
        _replace_paragraph(document, "Este informe", (
            f"Este informe estad\u00edstico presenta el uso de la plataforma Open LMS para el programa de "
            f"{program} durante el periodo {period}. Fue elaborado por el equipo del Campus Virtual."
        ))
        _replace_paragraph(document, "Para la", (
            f"Para la elaboraci\u00f3n del informe se analizaron {_number(total_events)} registros del sistema "
            "de seguimiento del Campus Virtual. La informaci\u00f3n fue consolidada autom\u00e1ticamente en Excel."
        ))
        sheet_descriptions = {
            "Original": "Datos originales y campos de fecha preparados.",
            "Tabla Dinamica Docentes": "Actividad docente consolidada por curso y mes.",
            "Docentes DG": "Representación gráfica de la actividad docente mensual.",
            "Tabla Dinamica Estudiantes": "Días activos y estudiantes únicos por curso y mes.",
            "Estudiantes DG": "Representación gráfica de los días de actividad estudiantil.",
            "Estudiantes DG2": "Promedio mensual de estudiantes que usaron el Campus Virtual.",
            "Tabla Dinamica Actividades": "Cantidad de acciones registradas en cada curso.",
            "Resumen Informe": "Indicadores y datos consolidados para el informe institucional.",
            "Diseño de Cursos": "Cantidad de cursos con contenido y sin contenido.",
        }
        sheet_lines = [
            f"{index}. Hoja «{name}»: {sheet_descriptions.get(name, 'Información complementaria del reporte.')}"
            for index, name in enumerate(data["sheet_names"], 1)
        ]
        sheet_lines.extend([""] * (10 - len(sheet_lines)))
        sheet_intro = (
            f"El archivo Excel adjunto contiene {len(data['sheet_names'])} hojas con el detalle y "
            f"los resultados del programa de {program} para el periodo {period}."
        )
        if _replace_marker(document, "{{DESCRIPCION_HOJAS}}", sheet_intro):
            for index, line in enumerate(sheet_lines[:10], 1):
                _replace_marker(document, f"{{{{HOJA_{index}}}}}", line)
        else:
            _replace_paragraph_group(document, "Este informe podr", [
                sheet_intro, *sheet_lines[:10],
            ])
        if data["monthly"]:
            peak = max(data["monthly"], key=lambda row: (row[1] or 0) + (row[3] or 0))
            monthly_text = (
                f"La mayor actividad se registr\u00f3 en {MONTH_NAMES.get(str(peak[0]).upper(), peak[0])}, "
                f"con {_number((peak[1] or 0) + (peak[3] or 0))} eventos combinados. La tabla compara "
                "el volumen de eventos con la cantidad de estudiantes activos por mes."
            )
            if not _replace_marker(document, "{{ANALISIS_MENSUAL}}", monthly_text):
                _replace_paragraph(document, "La actividad se concentr\u00f3", monthly_text)
        if teacher_rows:
            teacher_details = [
                f"{row[1]} en {row[0]}, con {row[2]} días activos durante el periodo."
                for row in teacher_rows[:3]
            ]
            teacher_details.extend([""] * (3 - len(teacher_details)))
            teacher_intro = (
                f"Los docentes con mayor continuidad se identificaron mediante los d\u00edas diferentes con "
                f"actividad. El primer lugar corresponde a {teacher_rows[0][1]} en {teacher_rows[0][0]}, "
                f"con {teacher_rows[0][2]} d\u00edas activos."
            )
            if _replace_marker(document, "{{ANALISIS_DOCENTES}}", teacher_intro):
                for index, detail in enumerate(teacher_details, 1):
                    _replace_marker(document, f"{{{{DOCENTE_{index}}}}}", detail)
            else:
                _replace_paragraph_group(document, "Con un uso que va", [teacher_intro, *teacher_details])
        if course_rows:
            course_details = [
                f"{row[0]}: {row[1]} eventos, {row[2]} estudiantes únicos y {row[3]} días activos."
                for row in course_rows[:3]
            ]
            course_details.extend([""] * (3 - len(course_details)))
            course_intro = (
                f"El curso con mayor continuidad estudiantil fue {course_rows[0][0]}, con "
                f"{course_rows[0][3]} d\u00edas activos, {course_rows[0][1]} eventos y "
                f"{course_rows[0][2]} estudiantes \u00fanicos."
            )
            if _replace_marker(document, "{{ANALISIS_ESTUDIANTES}}", course_intro):
                for index, detail in enumerate(course_details, 1):
                    _replace_marker(document, f"{{{{CURSO_{index}}}}}", detail)
            else:
                _replace_paragraph_group(document, "Con un acceso en promedio", [
                    course_intro, "Los cursos destacados son:", *course_details,
                ])
        design_text = (
            f"El programa de {program} cuenta con {design_total} cursos evaluados: {with_content} con "
            f"contenido y {data['design'].get('Sin contenido', 0)} sin contenido. Esto equivale a "
            f"{_number(content_percent, 1)} % de cursos con contenido."
        )
        if not _replace_marker(document, "{{ANALISIS_DISENO}}", design_text):
            _replace_containing(document, "cuenta con un total de 59 cursos", design_text)
        chart_titles = (
            "Eventos mensuales por tipo de usuario",
            "Días al mes de uso del Campus Virtual por parte de los docentes "
            f"de la facultad de {program} {period}",
            "Días del mes de uso del Campus Virtual por parte de los estudiantes "
            f"de la facultad de {program} {period}",
            "Promedio de estudiantes que usaron el Campus Virtual de la facultad "
            f"de {program} {period}",
            "Cursos con mayor continuidad estudiantil",
            "Diseño de Cursos",
        )
        for index, title in enumerate(chart_titles, 1):
            marker = f"{{{{TITULO_GRAFICA_{index}}}}}"
            if not _replace_marker(document, marker, title):
                _replace_paragraph(
                    document, f"Ilustración {index}",
                    f"Ilustración {index}. {title}.",
                )
        figure_captions = (
            f"Eventos mensuales por tipo de usuario {period}",
            f"Días al mes de uso de la plataforma Open LMS - Docentes {period}",
            f"Días al mes de uso de la plataforma Open LMS - Estudiantes {period}",
            f"Promedio de uso de la plataforma Open LMS - Estudiantes {period}",
            f"Cursos con mayor continuidad estudiantil {period}",
            f"Porcentaje de cursos con contenido {period}",
        )
        for index, caption in enumerate(figure_captions, 1):
            _set_figure_caption(document, index, caption)
        _style_figure_captions(document)
        _ensure_illustrations_cache(document, figure_captions)
        for paragraph in document.paragraphs:
            if "2026-1" in paragraph.text and period != "2026-1":
                paragraph.text = paragraph.text.replace("2026-1", period)
            updated_text = re.sub(
                r"Ingenier.a de Sistemas", program, paragraph.text,
                flags=re.IGNORECASE,
            )
            if updated_text != paragraph.text:
                paragraph.text = updated_text

        output_path.parent.mkdir(parents=True, exist_ok=True)
        temporary_output = output_path.with_name(output_path.stem + "_EN_PROCESO.docx")
        document.save(temporary_output)
        temporary_output.replace(output_path)
    return output_path
